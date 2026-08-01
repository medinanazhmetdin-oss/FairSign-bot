import os
import re
import io
import threading

import telebot
from flask import Flask
from dotenv import load_dotenv
from google import genai
from google.genai import types

# =========================
# НАСТРОЙКИ
# =========================

load_dotenv()

TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

if not TELEGRAM_TOKEN:
    raise RuntimeError("Не задана переменная TELEGRAM_TOKEN (проверь .env)")
if not GEMINI_API_KEY:
    raise RuntimeError("Не задана переменная GEMINI_API_KEY (проверь .env)")

GEMINI_MODEL = "gemini-2.5-flash"

MAX_CONTRACT_CHARS = 15000

# =========================
# ПОДКЛЮЧЕНИЕ
# =========================

bot = telebot.TeleBot(TELEGRAM_TOKEN)
ai = genai.Client(api_key=GEMINI_API_KEY)


# =========================
# ФИЛЬТР МАТА
# =========================

_PROFANITY_ROOTS = [
    r"бля[дт]",
    r"хуй|хуе|хуё",
    r"пизд",
    r"еба[тн]|ёба[тн]|ебал|ёбан",
    r"муда[кч]",
    r"сука|суч[ае]",
    r"гандон",
    r"долбоеб|долбоёб",
]
_PROFANITY_PATTERN = re.compile(
    r"\b(" + "|".join(_PROFANITY_ROOTS) + r")[а-яё]*\b",
    re.IGNORECASE,
)


def censor_profanity(text: str):
    found = False

    def _mask(match):
        nonlocal found
        found = True
        word = match.group(0)
        return word[0] + "*" * (len(word) - 1)

    new_text = _PROFANITY_PATTERN.sub(_mask, text)
    return new_text, found


# =========================
# ПРИВЕТСТВИЯ (не блокируем)
# =========================

_GREETING_PATTERN = re.compile(
    r"^\s*(привет|здравствуй|добрый день|добрый вечер|доброе утро|hi|hello)\W*\s*$",
    re.IGNORECASE,
)


def is_greeting(text: str) -> bool:
    return bool(_GREETING_PATTERN.match(text))


# =========================
# ИЗВЛЕЧЕНИЕ ТЕКСТА ИЗ ФАЙЛОВ
# =========================

def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


# =========================
# АНАЛИЗ ДОГОВОРА ЧЕРЕЗ GEMINI
# =========================

ANALYSIS_SYSTEM_PROMPT = """\
Ты — юридический ассистент по трудовому праву Республики Казахстан. Тебе дают
текст трудового договора. Проанализируй его МАКСИМАЛЬНО ТОЧНО и ПРЕДМЕТНО,
руководствуясь Трудовым кодексом РК (ТК РК). Не делай общих рассуждений без
опоры на конкретный текст договора — каждый вывод должен ссылаться на
конкретный пункт/формулировку из присланного текста.

Верни структурированный ответ на русском языке в следующем формате (используй
именно эти заголовки):

1. ОБЩИЕ СВЕДЕНИЯ — тип договора (срочный/бессрочный), стороны, дата, срок
   действия. Указывай ТОЛЬКО то, что реально написано в тексте; если пункт
   не найден в договоре — прямо пиши "не указано в тексте", не придумывай.

2. КЛЮЧЕВЫЕ УСЛОВИЯ — должность, размер и порядок оплаты труда, режим и место
   работы, испытательный срок. По каждому пункту — короткая цитата или
   пересказ формулировки из договора, из которой это следует.

3. РИСКИ И ПРОБЛЕМНЫЕ ПУНКТЫ — это САМЫЙ ВАЖНЫЙ раздел, разбери его подробно.
   Для КАЖДОЙ найденной проблемы дай отдельный блок строго по такой структуре:

   ⚠️ [Название проблемы]
   - Что в договоре: точная формулировка/пункт из текста.
   - Почему это плохо: конкретное нарушение или риск, со ссылкой на статью
     ТК РК (например: ст. 23 — дискриминация, ст. 24/28 — обязательные
     условия договора, ст. 36 — испытательный срок, ст. 65 — рабочее время,
     ст. 71-72 — сверхурочная работа, ст. 88 — отпуска, ст. 131-132 —
     расторжение договора, ст. 113 — материальная ответственность и т.д.).
   - Уровень критичности: 🔴 критично / 🟠 средне / 🟡 незначительно.
   - Совет: конкретная рекомендация, что именно исправить.

   Если проблем не найдено — прямо напиши "существенных нарушений не
   обнаружено", не придумывай риски искусственно.

4. ЧЕГО НЕ ХВАТАЕТ — обязательные по ст. 28 ТК РК условия, которые
   отсутствуют в тексте. По каждому пункту — конкретный совет, какую
   формулировку добавить в договор.

5. ИТОГОВАЯ ОЦЕНКА — можно ли подписывать договор как есть, что обязательно
   нужно исправить перед подписанием, и общая оценка риска (низкий/средний/
   высокий).

Требования к точности:
- Не приписывай договору то, чего в нём нет.
- Если текст обрезан или неполный — упомяни об этом в итоговой оценке.
- Если переданный текст вообще не похож на трудовой договор — прямо напиши
  об этом и не придумывай анализ.
- Отметь, что это предварительный анализ и не заменяет консультацию юриста в РК.
"""


def analyze_contract(contract_text: str) -> str:
    contract_text = contract_text[:MAX_CONTRACT_CHARS]

    response = ai.models.generate_content(
        model=GEMINI_MODEL,
        contents=f"Вот текст трудового договора:\n\n{contract_text}",
        config=types.GenerateContentConfig(
            system_instruction=ANALYSIS_SYSTEM_PROMPT,
            temperature=0.2,
            max_output_tokens=3500,
        ),
    )
    if not response.text:
        raise RuntimeError("Gemini вернул пустой ответ")
    return response.text


def send_long_message(chat_id, text: str):
    limit = 4000
    for i in range(0, len(text), limit):
        bot.send_message(chat_id, text[i : i + limit])


# =========================
# START
# =========================

@bot.message_handler(commands=["start"])
def start(message):
    bot.send_message(
        message.chat.id,
        "Привет! 🤖 Я FairSign — бот для анализа трудовых договоров по "
        "законодательству Республики Казахстан (ТК РК).\n\n"
        "Пришли мне текст договора или файл (.docx / .pdf / .txt), и я разберу "
        "его: условия, риски со ссылками на статьи ТК РК, чего не хватает, "
        "и итоговую оценку.\n\n"
        "Ненормативная лексика в сообщениях будет автоматически скрыта."
    )


# =========================
# ФАЙЛЫ (.docx / .pdf / .txt)
# =========================

@bot.message_handler(content_types=["document"])
def handle_document(message):
    caption = message.caption or ""
    if caption:
        clean_caption, had_profanity = censor_profanity(caption)
        if had_profanity:
            bot.send_message(
                message.chat.id,
                f"⚠️ Подпись к файлу содержала ненормативную лексику, скрыто:\n\n{clean_caption}",
            )

    file_name = (message.document.file_name or "").lower()
    file_info = bot.get_file(message.document.file_id)
    file_bytes = bot.download_file(file_info.file_path)

    try:
        if file_name.endswith(".docx"):
            text = extract_text_from_docx(file_bytes)
        elif file_name.endswith(".pdf"):
            text = extract_text_from_pdf(file_bytes)
        elif file_name.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            bot.send_message(message.chat.id, "Поддерживаются только файлы .docx, .pdf и .txt.")
            return
    except Exception as error:
        print("Ошибка чтения файла:", error)
        bot.send_message(message.chat.id, f"Не удалось прочитать файл: {error}")
        return

    if len(text.strip()) < 100:
        bot.send_message(
            message.chat.id,
            "Не удалось извлечь достаточно текста из файла (возможно, это скан без текстового слоя).",
        )
        return

    clean_text, _ = censor_profanity(text)

    bot.send_chat_action(message.chat.id, "typing")
    bot.send_message(message.chat.id, "Анализирую договор через Gemini, это может занять полминуты...")

    try:
        result = analyze_contract(clean_text)
    except Exception as error:
        print("Ошибка Gemini:", error)
        bot.send_message(message.chat.id, f"❌ Не удалось выполнить анализ:\n\n{error}")
        return

    send_long_message(message.chat.id, result)


# =========================
# ТЕКСТОВЫЕ СООБЩЕНИЯ
# =========================

@bot.message_handler(func=lambda message: True)
def handle_text(message):
    raw_text = message.text or ""

    if is_greeting(raw_text):
        bot.send_message(
            message.chat.id,
            "Здравствуйте! Пришлите текст трудового договора или файл, чтобы я его проанализировал.",
        )
        return

    clean_text, had_profanity = censor_profanity(raw_text)
    if had_profanity:
        bot.send_message(
            message.chat.id,
            f"⚠️ Сообщение содержит ненормативную лексику, она была скрыта:\n\n{clean_text}",
        )

    if len(clean_text.strip()) < 100:
        try:
            bot.send_chat_action(message.chat.id, "typing")
            response = ai.models.generate_content(
                model=GEMINI_MODEL,
                contents=clean_text,
            )
            bot.send_message(message.chat.id, response.text)
        except Exception as error:
            print("Ошибка Gemini:", error)
            bot.send_message(message.chat.id, f"❌ Ошибка Gemini:\n\n{error}")
        return

    bot.send_chat_action(message.chat.id, "typing")
    bot.send_message(message.chat.id, "Анализирую договор через Gemini, это может занять полминуты...")

    try:
        result = analyze_contract(clean_text)
    except Exception as error:
        print("Ошибка Gemini:", error)
        bot.send_message(message.chat.id, f"❌ Не удалось выполнить анализ:\n\n{error}")
        return

    send_long_message(message.chat.id, result)


# =========================
# KEEP-ALIVE СЕРВЕР (для Render — обязателен)
# =========================

keep_alive_app = Flask(__name__)


@keep_alive_app.route("/")
def home():
    return "FairSign bot is running!"


def run_keep_alive():
    port = int(os.environ.get("PORT", 3000))
    keep_alive_app.run(host="0.0.0.0", port=port)


def start_keep_alive():
    t = threading.Thread(target=run_keep_alive)
    t.daemon = True
    t.start()


# =========================
# ЗАПУСК
# =========================

start_keep_alive()
print("Бот запущен!")
bot.infinity_polling()
